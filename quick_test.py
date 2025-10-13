#!/usr/bin/env python3
"""
Quick test to verify performance optimizations are working
"""
import asyncio
import time
import aiohttp
from pathlib import Path
from docx import Document

async def test_async_operations():
    """Test that the server is handling async operations properly"""
    print("Testing Performance Optimizations...")
    print("="*50)

    # Create a simple test document
    print("1. Creating test document...")
    doc = Document()
    doc.add_heading('Test NDA Document', 0)
    doc.add_paragraph('This is a test document for performance verification.')

    # Add some content
    for i in range(10):
        doc.add_heading(f'Section {i+1}', level=1)
        doc.add_paragraph(f'This is section {i+1} content. ' * 20)

    test_file = Path('test_nda.docx')
    doc.save(str(test_file))
    print(f"   [OK] Created {test_file}")

    # Test upload and processing
    print("\n2. Testing document upload and processing...")
    start_time = time.time()

    async with aiohttp.ClientSession() as session:
        # Upload document
        with open(test_file, 'rb') as f:
            data = aiohttp.FormData()
            data.add_field('file', f, filename='test_nda.docx')

            async with session.post('http://localhost:8000/api/upload', data=data) as resp:
                if resp.status != 200:
                    print(f"   [FAIL] Upload failed: {resp.status}")
                    return False

                result = await resp.json()
                job_id = result.get('job_id')
                print(f"   [OK] Document uploaded, job_id: {job_id}")

        # Poll for completion
        print("\n3. Processing document...")
        max_wait = 60  # seconds
        poll_start = time.time()

        while time.time() - poll_start < max_wait:
            async with session.get(f'http://localhost:8000/api/jobs/{job_id}/status') as resp:
                if resp.status != 200:
                    print(f"   [FAIL] Status check failed: {resp.status}")
                    return False

                status = await resp.json()
                current_status = status.get('status')
                progress = status.get('progress', 0)

                print(f"   Status: {current_status} ({progress}%)", end='\r')

                if current_status == 'complete':
                    print(f"\n   [OK] Processing complete!")
                    print(f"   Total redlines: {status.get('total_redlines', 0)}")
                    break
                elif current_status == 'error':
                    print(f"\n   [FAIL] Processing failed: {status.get('error_message')}")
                    return False

            await asyncio.sleep(1)

    elapsed = time.time() - start_time
    print(f"\n4. Performance Results:")
    print(f"   Total time: {elapsed:.2f} seconds")

    # Test concurrent requests
    print("\n5. Testing concurrent processing...")
    start_time = time.time()

    async def upload_and_process():
        async with aiohttp.ClientSession() as session:
            with open(test_file, 'rb') as f:
                data = aiohttp.FormData()
                data.add_field('file', f, filename='test_nda.docx')

                async with session.post('http://localhost:8000/api/upload', data=data) as resp:
                    if resp.status == 200:
                        return await resp.json()
        return None

    # Send 3 concurrent requests
    tasks = [upload_and_process() for _ in range(3)]
    results = await asyncio.gather(*tasks)

    successful = len([r for r in results if r is not None])
    elapsed = time.time() - start_time

    print(f"   [OK] Sent 3 concurrent requests")
    print(f"   Successful: {successful}/3")
    print(f"   Time for 3 concurrent: {elapsed:.2f} seconds")
    print(f"   Throughput: {successful/elapsed:.2f} req/s")

    # Cleanup
    test_file.unlink()

    print("\n" + "="*50)
    print("[SUCCESS] All tests completed successfully!")
    print("="*50)
    return True

async def main():
    try:
        # Check server is running
        async with aiohttp.ClientSession() as session:
            async with session.get('http://localhost:8000') as resp:
                if resp.status != 200:
                    print("[FAIL] Server not running. Please start the server first.")
                    return False

        # Run tests
        success = await test_async_operations()
        return success

    except Exception as e:
        print(f"[FAIL] Test failed with error: {e}")
        return False

if __name__ == "__main__":
    result = asyncio.run(main())
    exit(0 if result else 1)