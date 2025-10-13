#!/usr/bin/env python3
"""
Performance Benchmark Script for NDA Redlining System
Measures the impact of optimizations:
- Async LLM clients
- Parallel execution
- Non-blocking I/O
- Memory optimizations
"""

import asyncio
import time
import json
import os
import sys
import psutil
import statistics
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Tuple
import aiohttp
import logging
from docx import Document
from docx.shared import Pt
import random

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# API Configuration
API_BASE_URL = "http://localhost:8000"
UPLOAD_ENDPOINT = f"{API_BASE_URL}/api/upload"
STATUS_ENDPOINT = f"{API_BASE_URL}/api/jobs"

class PerformanceBenchmark:
    """Comprehensive performance benchmark suite"""

    def __init__(self):
        self.results = {
            'timestamp': datetime.now().isoformat(),
            'system_info': self._get_system_info(),
            'benchmarks': {}
        }
        self.test_docs_path = Path("test_docs")
        self.test_docs_path.mkdir(exist_ok=True)

    def _get_system_info(self) -> Dict:
        """Get system information"""
        return {
            'cpu_count': psutil.cpu_count(),
            'memory_gb': round(psutil.virtual_memory().total / (1024**3), 2),
            'python_version': sys.version,
            'platform': sys.platform
        }

    def generate_test_document(self, pages: int, filename: str) -> Path:
        """Generate test NDA document with specified pages"""
        doc = Document()

        doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)
        doc.add_paragraph(f'This Agreement dated as of {datetime.now().date()}')

        # Calculate content: ~3 paragraphs per page
        num_sections = pages * 3

        sections = [
            ("WHEREAS", "Party A desires to disclose certain confidential information..."),
            ("Definitions", "Confidential Information means any data or information..."),
            ("Obligations", "Receiving Party shall hold in confidence and not disclose..."),
            ("Permitted Disclosures", "Receiving Party may disclose Confidential Information..."),
            ("Term", "This Agreement shall commence on the Effective Date..."),
            ("Return of Materials", "Upon termination, Receiving Party shall return..."),
            ("Remedies", "Receiving Party acknowledges that breach may cause irreparable harm..."),
            ("Miscellaneous", "This Agreement constitutes the entire agreement...")
        ]

        for i in range(num_sections):
            section_title, section_text = sections[i % len(sections)]
            doc.add_heading(f'{i+1}. {section_title}', level=1)

            # Add 2-4 paragraphs per section
            for _ in range(random.randint(2, 4)):
                para = doc.add_paragraph()
                para.add_run(section_text + " " + ("Additional legal language. " * 15))

            # Add table every 10 sections for table processing test
            if (i + 1) % 10 == 0:
                table = doc.add_table(rows=5, cols=3)
                table.style = 'Table Grid'
                for row_idx in range(5):
                    for col_idx in range(3):
                        table.cell(row_idx, col_idx).text = f"Cell R{row_idx}C{col_idx}"

        filepath = self.test_docs_path / filename
        doc.save(str(filepath))
        logger.info(f"Generated test document: {filename} (~{pages} pages)")
        return filepath

    async def measure_single_request(self, filepath: Path) -> Dict:
        """Measure performance for a single document processing request"""
        start_time = time.perf_counter()
        start_memory = psutil.Process().memory_info().rss / 1024 / 1024  # MB

        async with aiohttp.ClientSession() as session:
            # Upload document
            with open(filepath, 'rb') as f:
                data = aiohttp.FormData()
                data.add_field('file', f, filename=filepath.name)

                async with session.post(UPLOAD_ENDPOINT, data=data) as resp:
                    upload_result = await resp.json()
                    job_id = upload_result['job_id']

            # Poll for completion
            while True:
                async with session.get(f"{STATUS_ENDPOINT}/{job_id}/status") as resp:
                    status = await resp.json()

                    if status['status'] in ['complete', 'error']:
                        break

                await asyncio.sleep(1)

        end_time = time.perf_counter()
        end_memory = psutil.Process().memory_info().rss / 1024 / 1024  # MB

        return {
            'duration_s': round(end_time - start_time, 2),
            'memory_delta_mb': round(end_memory - start_memory, 2),
            'status': status['status'],
            'redlines': status.get('total_redlines', 0)
        }

    async def measure_concurrent_requests(self, filepath: Path, concurrent: int) -> Dict:
        """Measure performance for concurrent requests"""
        start_time = time.perf_counter()

        tasks = []
        for _ in range(concurrent):
            tasks.append(self.measure_single_request(filepath))

        results = await asyncio.gather(*tasks, return_exceptions=True)

        end_time = time.perf_counter()
        total_duration = end_time - start_time

        successful = [r for r in results if isinstance(r, dict) and r['status'] == 'complete']
        failed = len(results) - len(successful)

        durations = [r['duration_s'] for r in successful]

        return {
            'total_duration_s': round(total_duration, 2),
            'concurrent_requests': concurrent,
            'successful': len(successful),
            'failed': failed,
            'throughput_req_per_s': round(len(successful) / total_duration, 3) if total_duration > 0 else 0,
            'avg_duration_s': round(statistics.mean(durations), 2) if durations else 0,
            'min_duration_s': round(min(durations), 2) if durations else 0,
            'max_duration_s': round(max(durations), 2) if durations else 0,
            'speedup_factor': round(statistics.mean(durations) / total_duration * concurrent, 2) if durations and total_duration > 0 else 0
        }

    async def run_benchmark_suite(self) -> None:
        """Run complete benchmark suite"""
        logger.info("="*60)
        logger.info("NDA Redlining Performance Benchmark")
        logger.info("="*60)

        # Generate test documents
        test_cases = [
            (5, "nda_5pg.docx"),
            (25, "nda_25pg.docx"),
            (100, "nda_100pg.docx")
        ]

        for pages, filename in test_cases:
            if not (self.test_docs_path / filename).exists():
                self.generate_test_document(pages, filename)

        # Benchmark 1: Single request latency
        logger.info("\n=== Benchmark 1: Single Request Latency ===")
        for pages, filename in test_cases:
            filepath = self.test_docs_path / filename
            logger.info(f"Testing {filename}...")

            # Run 3 iterations and take average
            iterations = []
            for i in range(3):
                logger.info(f"  Iteration {i+1}/3...")
                result = await self.measure_single_request(filepath)
                iterations.append(result)
                await asyncio.sleep(2)  # Cool down between tests

            avg_duration = statistics.mean([r['duration_s'] for r in iterations])
            avg_memory = statistics.mean([r['memory_delta_mb'] for r in iterations])

            self.results['benchmarks'][f'single_{pages}pg'] = {
                'pages': pages,
                'avg_duration_s': round(avg_duration, 2),
                'avg_memory_delta_mb': round(avg_memory, 2),
                'iterations': iterations
            }

            logger.info(f"  Average: {avg_duration:.2f}s, Memory: {avg_memory:.2f}MB")

        # Benchmark 2: Concurrent request handling
        logger.info("\n=== Benchmark 2: Concurrent Request Handling ===")
        concurrent_tests = [5, 10]
        test_file = self.test_docs_path / "nda_25pg.docx"

        for concurrent in concurrent_tests:
            logger.info(f"Testing {concurrent} concurrent requests...")
            result = await self.measure_concurrent_requests(test_file, concurrent)

            self.results['benchmarks'][f'concurrent_{concurrent}'] = result

            logger.info(f"  Total time: {result['total_duration_s']}s")
            logger.info(f"  Throughput: {result['throughput_req_per_s']} req/s")
            logger.info(f"  Speedup: {result['speedup_factor']}×")

            await asyncio.sleep(5)  # Cool down between tests

        # Benchmark 3: Memory leak detection
        logger.info("\n=== Benchmark 3: Memory Leak Detection ===")
        test_file = self.test_docs_path / "nda_25pg.docx"
        memory_samples = []

        for i in range(10):
            logger.info(f"  Iteration {i+1}/10...")
            before_memory = psutil.Process().memory_info().rss / 1024 / 1024

            await self.measure_single_request(test_file)

            after_memory = psutil.Process().memory_info().rss / 1024 / 1024
            memory_samples.append(after_memory - before_memory)

            await asyncio.sleep(2)

        avg_memory_growth = statistics.mean(memory_samples)
        self.results['benchmarks']['memory_leak_test'] = {
            'iterations': 10,
            'avg_memory_growth_mb': round(avg_memory_growth, 2),
            'max_memory_growth_mb': round(max(memory_samples), 2),
            'samples': [round(s, 2) for s in memory_samples]
        }

        logger.info(f"  Average memory growth: {avg_memory_growth:.2f}MB/iteration")

        # Save results
        self._save_results()
        self._print_summary()

    def _save_results(self):
        """Save benchmark results to JSON file"""
        filename = f"benchmark_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(filename, 'w') as f:
            json.dump(self.results, f, indent=2)
        logger.info(f"\nResults saved to: {filename}")

    def _print_summary(self):
        """Print benchmark summary"""
        logger.info("\n" + "="*60)
        logger.info("BENCHMARK SUMMARY")
        logger.info("="*60)

        # Single request performance
        if 'single_25pg' in self.results['benchmarks']:
            result = self.results['benchmarks']['single_25pg']
            logger.info(f"\n25-page document processing:")
            logger.info(f"  Average time: {result['avg_duration_s']}s")
            logger.info(f"  Memory usage: {result['avg_memory_delta_mb']}MB")

        # Concurrent performance
        if 'concurrent_10' in self.results['benchmarks']:
            result = self.results['benchmarks']['concurrent_10']
            logger.info(f"\n10 concurrent requests:")
            logger.info(f"  Total time: {result['total_duration_s']}s")
            logger.info(f"  Throughput: {result['throughput_req_per_s']} req/s")
            logger.info(f"  Speedup factor: {result['speedup_factor']}×")

        # Memory leak
        if 'memory_leak_test' in self.results['benchmarks']:
            result = self.results['benchmarks']['memory_leak_test']
            logger.info(f"\nMemory leak test (10 iterations):")
            logger.info(f"  Average growth: {result['avg_memory_growth_mb']}MB/iter")

            if result['avg_memory_growth_mb'] > 10:
                logger.warning("  ⚠️ POTENTIAL MEMORY LEAK DETECTED")
            else:
                logger.info("  ✓ No significant memory leak detected")

        logger.info("\n" + "="*60)
        logger.info("Benchmark complete!")
        logger.info("="*60)

    def compare_with_baseline(self, baseline_file: str):
        """Compare current results with baseline"""
        try:
            with open(baseline_file, 'r') as f:
                baseline = json.load(f)

            logger.info("\n" + "="*60)
            logger.info("PERFORMANCE COMPARISON WITH BASELINE")
            logger.info("="*60)

            # Compare single request performance
            for test_case in ['single_5pg', 'single_25pg', 'single_100pg']:
                if test_case in self.results['benchmarks'] and test_case in baseline['benchmarks']:
                    current = self.results['benchmarks'][test_case]['avg_duration_s']
                    base = baseline['benchmarks'][test_case]['avg_duration_s']
                    speedup = base / current if current > 0 else 0

                    logger.info(f"\n{test_case}:")
                    logger.info(f"  Baseline: {base}s")
                    logger.info(f"  Current: {current}s")
                    logger.info(f"  Speedup: {speedup:.2f}× {'🎉' if speedup > 1.5 else ''}")

            # Compare concurrent performance
            for test_case in ['concurrent_5', 'concurrent_10']:
                if test_case in self.results['benchmarks'] and test_case in baseline['benchmarks']:
                    current_throughput = self.results['benchmarks'][test_case]['throughput_req_per_s']
                    base_throughput = baseline['benchmarks'][test_case]['throughput_req_per_s']
                    improvement = (current_throughput / base_throughput - 1) * 100 if base_throughput > 0 else 0

                    logger.info(f"\n{test_case} throughput:")
                    logger.info(f"  Baseline: {base_throughput} req/s")
                    logger.info(f"  Current: {current_throughput} req/s")
                    logger.info(f"  Improvement: {improvement:+.1f}% {'🎉' if improvement > 50 else ''}")

        except FileNotFoundError:
            logger.warning(f"Baseline file not found: {baseline_file}")
        except Exception as e:
            logger.error(f"Error comparing with baseline: {e}")


async def main():
    """Run the benchmark suite"""
    benchmark = PerformanceBenchmark()

    # Check if server is running
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(API_BASE_URL) as resp:
                if resp.status != 200:
                    logger.error("Server not responding. Please start the server first.")
                    return
    except Exception as e:
        logger.error(f"Cannot connect to server at {API_BASE_URL}")
        logger.error("Please start the server with: python -m uvicorn backend.app.main:app --reload")
        return

    # Run benchmarks
    await benchmark.run_benchmark_suite()

    # Compare with baseline if exists
    baseline_files = list(Path('.').glob('baseline_*.json'))
    if baseline_files:
        latest_baseline = max(baseline_files, key=lambda p: p.stat().st_mtime)
        benchmark.compare_with_baseline(str(latest_baseline))


if __name__ == "__main__":
    asyncio.run(main())