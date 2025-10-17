"""Create a comprehensive test NDA document with various clause types"""
from docx import Document

# Create a new Document
doc = Document()

# Add title
doc.add_heading('NON-DISCLOSURE AGREEMENT', 0)

# Add date and parties
doc.add_paragraph('This Non-Disclosure Agreement ("Agreement") is entered into as of January 1, 2025.')
doc.add_paragraph()

doc.add_paragraph('BETWEEN:')
doc.add_paragraph('ABC Corporation, a company incorporated under the laws of Delaware ("Disclosing Party")')
doc.add_paragraph()
doc.add_paragraph('AND:')
doc.add_paragraph('XYZ Inc., a company incorporated under the laws of California ("Receiving Party")')
doc.add_paragraph()

# Add problematic clauses that should trigger redlines
doc.add_heading('1. Definition of Confidential Information', level=1)
doc.add_paragraph(
    'Confidential Information means any and all information disclosed by the Disclosing Party '
    'to the Receiving Party, whether orally or in writing, that is designated as confidential '
    'or that reasonably should be understood to be confidential given the nature of the information '
    'and the circumstances of disclosure.'
)
doc.add_paragraph()

doc.add_heading('2. Obligations', level=1)
doc.add_paragraph(
    'The Receiving Party agrees to hold the Confidential Information in strict confidence '
    'and shall not disclose such information to any third parties without the prior written '
    'consent of the Disclosing Party. The Receiving Party shall indemnify and hold harmless '
    'the Disclosing Party from any and all claims, damages, losses, and expenses arising out of '
    'or relating to any breach of this Agreement.'
)
doc.add_paragraph()

doc.add_heading('3. Term and Termination', level=1)
doc.add_paragraph(
    'This Agreement shall remain in effect in perpetuity and the obligations hereunder shall '
    'survive indefinitely. The Receiving Party\'s obligations shall continue even after the '
    'termination of any business relationship between the parties.'
)
doc.add_paragraph()

doc.add_heading('4. Non-Compete', level=1)
doc.add_paragraph(
    'The Receiving Party agrees that during the term of this Agreement and for a period of '
    '10 years thereafter, it shall not engage in any business that competes with the Disclosing '
    'Party anywhere in the world.'
)
doc.add_paragraph()

doc.add_heading('5. Governing Law', level=1)
doc.add_paragraph(
    'This Agreement shall be governed by and construed in accordance with the laws of '
    'the Cayman Islands, without regard to its conflict of law provisions.'
)
doc.add_paragraph()

doc.add_heading('6. Dispute Resolution', level=1)
doc.add_paragraph(
    'Any disputes arising under this Agreement shall be resolved exclusively by binding '
    'arbitration in the Cayman Islands. The prevailing party shall be entitled to recover '
    'all attorneys\' fees and costs.'
)
doc.add_paragraph()

doc.add_heading('7. Entire Agreement', level=1)
doc.add_paragraph(
    'This Agreement constitutes the entire agreement between the parties and supersedes all '
    'prior agreements and understandings, whether written or oral.'
)

# Save the document
doc.save('test_comprehensive_nda.docx')
print('Created test_comprehensive_nda.docx with multiple clause types')
print('  - Includes overly broad confidentiality definition')
print('  - Includes unlimited indemnification')
print('  - Includes perpetual term')
print('  - Includes unreasonable non-compete (10 years, worldwide)')
print('  - Uses unfavorable jurisdiction (Cayman Islands)')
