"""
WorkingTextIndexer - Bidirectional mapping between normalized text and DOCX structure
Critical for precise span alignment when applying track changes
"""
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
from typing import List, Dict, Tuple, Optional
import re


class TextMapping:
    """Represents a span in the working text and its DOCX location"""
    def __init__(self, start: int, end: int, p_idx: int, r_idx: int,
                 original: str, element_type: str = 'paragraph', table_info: Dict = None):
        self.start = start
        self.end = end
        self.p_idx = p_idx
        self.r_idx = r_idx
        self.original = original
        self.element_type = element_type
        self.table_info = table_info or {}


class WorkingTextIndexer:
    """
    Build a normalized text representation with bidirectional mappings to DOCX structure.

    Handles:
    - Paragraphs and runs
    - Tables and cells
    - Headers and footers
    - Text normalization (whitespace, special chars)
    - Run merging for adjacent identical formatting
    """

    def __init__(self):
        self.working_text = ""
        self.mappings: List[TextMapping] = []
        self.doc = None

    def normalize_text(self, text: str) -> str:
        """Normalize text for consistent processing"""
        if not text:
            return ""

        # Normalize whitespace but preserve structure
        text = re.sub(r'\s+', ' ', text)

        # Remove zero-width characters
        text = re.sub(r'[\u200b\ufeff]', '', text)

        return text

    def should_merge_runs(self, run1, run2) -> bool:
        """Check if two adjacent runs have identical formatting"""
        if not run1 or not run2:
            return False

        # Compare key formatting properties
        return (
            run1.bold == run2.bold and
            run1.italic == run2.italic and
            run1.underline == run2.underline and
            run1.font.name == run2.font.name and
            run1.font.size == run2.font.size
        )

    def normalize_runs(self, runs) -> List:
        """Merge adjacent runs with identical formatting"""
        if not runs:
            return []

        normalized = []
        current_run = None
        current_text = ""

        for run in runs:
            if current_run is None:
                current_run = run
                current_text = run.text
            elif self.should_merge_runs(current_run, run):
                # Merge with previous run
                current_text += run.text
            else:
                # Save previous and start new
                if current_text:
                    normalized.append({
                        'run': current_run,
                        'text': current_text
                    })
                current_run = run
                current_text = run.text

        # Don't forget the last run
        if current_run and current_text:
            normalized.append({
                'run': current_run,
                'text': current_text
            })

        return normalized

    def index_paragraph(self, paragraph: Paragraph, p_idx: int, element_type: str = 'paragraph', table_info: Dict = None):
        """Index a single paragraph with its runs"""
        normalized_runs = self.normalize_runs(paragraph.runs)

        for r_idx, run_data in enumerate(normalized_runs):
            run = run_data['run']
            text = run_data['text']

            start = len(self.working_text)
            normalized = self.normalize_text(text)
            self.working_text += normalized
            end = len(self.working_text)

            if normalized:  # Only add non-empty mappings
                mapping = TextMapping(
                    start=start,
                    end=end,
                    p_idx=p_idx,
                    r_idx=r_idx,
                    original=text,
                    element_type=element_type,
                    table_info=table_info
                )
                self.mappings.append(mapping)

        # Add paragraph break
        if normalized_runs:
            self.working_text += "\n"

    def index_table(self, table: Table, table_idx: int):
        """Index all cells in a table"""
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                table_info = {
                    'table_idx': table_idx,
                    'row': row_idx,
                    'col': col_idx
                }

                for p_idx, paragraph in enumerate(cell.paragraphs):
                    self.index_paragraph(
                        paragraph,
                        p_idx,
                        element_type='table_cell',
                        table_info=table_info
                    )

    def build_index(self, doc: Document):
        """Build complete index of document"""
        self.doc = doc
        self.working_text = ""
        self.mappings = []

        # Index main document paragraphs
        for p_idx, paragraph in enumerate(doc.paragraphs):
            self.index_paragraph(paragraph, p_idx)

        # Index tables
        for table_idx, table in enumerate(doc.tables):
            self.index_table(table, table_idx)

        # Index headers
        for section in doc.sections:
            if section.header:
                for p_idx, paragraph in enumerate(section.header.paragraphs):
                    self.index_paragraph(paragraph, p_idx, element_type='header')

            if section.footer:
                for p_idx, paragraph in enumerate(section.footer.paragraphs):
                    self.index_paragraph(paragraph, p_idx, element_type='footer')

    def find_spans(self, start: int, end: int) -> List[TextMapping]:
        """
        Find all DOCX mappings that overlap with the given text span.
        Binary search for efficiency.
        """
        result = []

        for mapping in self.mappings:
            # Check for overlap
            if mapping.end <= start:
                continue
            if mapping.start >= end:
                break

            result.append(mapping)

        return result

    def find_exact_span(self, text: str, start_hint: int = 0) -> Optional[Tuple[int, int]]:
        """Find exact match of text in working_text, starting from hint"""
        normalized = self.normalize_text(text)
        idx = self.working_text.find(normalized, start_hint)

        if idx != -1:
            return (idx, idx + len(normalized))

        return None

    def get_context(self, start: int, end: int, chars_before: int = 50, chars_after: int = 50) -> str:
        """Get surrounding context for a text span"""
        context_start = max(0, start - chars_before)
        context_end = min(len(self.working_text), end + chars_after)

        return self.working_text[context_start:context_end]

    def get_paragraph_and_run(self, mapping: TextMapping):
        """Get the actual paragraph and run objects from a mapping"""
        if mapping.element_type == 'paragraph':
            paragraph = self.doc.paragraphs[mapping.p_idx]
            normalized_runs = self.normalize_runs(paragraph.runs)
            if mapping.r_idx < len(normalized_runs):
                return paragraph, normalized_runs[mapping.r_idx]['run']

        elif mapping.element_type == 'table_cell':
            table = self.doc.tables[mapping.table_info['table_idx']]
            cell = table.rows[mapping.table_info['row']].cells[mapping.table_info['col']]
            paragraph = cell.paragraphs[mapping.p_idx]
            normalized_runs = self.normalize_runs(paragraph.runs)
            if mapping.r_idx < len(normalized_runs):
                return paragraph, normalized_runs[mapping.r_idx]['run']

        return None, None

    def debug_span(self, start: int, end: int):
        """Debug helper to inspect a text span and its mappings"""
        text = self.working_text[start:end]
        mappings = self.find_spans(start, end)

        print(f"Text span [{start}:{end}]: '{text}'")
        print(f"Found {len(mappings)} mapping(s):")

        for m in mappings:
            print(f"  - [{m.start}:{m.end}] p={m.p_idx} r={m.r_idx} type={m.element_type}")
            print(f"    Original: '{m.original}'")
