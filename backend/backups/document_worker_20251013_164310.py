"""
Document processing worker
Handles async NDA processing pipeline with Redis queue support
"""
import os
import asyncio
import json
import uuid
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Optional
from docx import Document

from ..core.text_indexer import WorkingTextIndexer
from ..core.rule_engine import RuleEngine
from ..core.llm_orchestrator import LLMOrchestrator
from ..core.docx_engine import TrackChangesEngine, RedlineValidator
from ..models.schemas import JobStatus, RedlineModel, RedlineSeverity, RedlineSource
from ..models.checklist_rules import get_rule_explanation
from .redis_job_queue import RedisJobQueue, JobPriority, JobStatus as RedisJobStatus


class DocumentProcessor:
    """
    Process a single NDA document through the pipeline:
    1. Parse DOCX and build index
    2. Apply deterministic rules
    3. LLM analysis with validation
    4. Generate redlined document
    """

    def __init__(self, storage_path: str = "./storage"):
        self.storage_path = Path(storage_path)
        self.rule_engine = RuleEngine()
        self.llm_orchestrator = LLMOrchestrator()

    async def process_document(
        self,
        job_id: str,
        file_path: str,
        status_callback=None
    ) -> Dict:
        """
        Process a document through the full pipeline.

        Args:
            job_id: Unique job identifier
            file_path: Path to uploaded DOCX
            status_callback: Async callback for status updates

        Returns:
            Dict with processing results
        """
        try:
            # Update status: parsing
            await self._update_status(status_callback, JobStatus.PARSING, 10)

            # 1. Parse DOCX
            doc = Document(file_path)
            indexer = WorkingTextIndexer()
            indexer.build_index(doc)

            working_text = indexer.working_text

            # Save working text for debugging
            working_text_path = self.storage_path / "working" / f"{job_id}.txt"
            working_text_path.parent.mkdir(parents=True, exist_ok=True)
            working_text_path.write_text(working_text, encoding='utf-8')

            # Update status: applying rules
            await self._update_status(status_callback, JobStatus.APPLYING_RULES, 30)

            # 2. Apply deterministic rules
            rule_redlines = self.rule_engine.apply_rules(working_text)

            print(f"Job {job_id}: Found {len(rule_redlines)} rule-based redlines")

            # Update status: analyzing
            await self._update_status(status_callback, JobStatus.ANALYZING, 50)

            # 3. LLM analysis (now async)
            llm_redlines = await self.llm_orchestrator.analyze(working_text, rule_redlines)

            print(f"Job {job_id}: Found {len(llm_redlines)} LLM redlines")

            # Combine all redlines
            all_redlines = rule_redlines + llm_redlines

            # Validate all redlines
            valid_redlines = RedlineValidator.validate_all(all_redlines, working_text)

            print(f"Job {job_id}: {len(valid_redlines)} valid redlines after validation")

            # Convert to RedlineModel objects
            redline_models = self._convert_to_models(valid_redlines)

            # Update status: generating
            await self._update_status(status_callback, JobStatus.GENERATING, 80)

            # 4. Generate redlined document
            output_path = await self._generate_redlined_doc(
                job_id,
                doc,
                indexer,
                valid_redlines
            )

            # Update status: complete
            await self._update_status(status_callback, JobStatus.COMPLETE, 100)

            # Return results
            result = {
                'job_id': job_id,
                'status': JobStatus.COMPLETE,
                'total_redlines': len(valid_redlines),
                'rule_redlines': len(rule_redlines),
                'llm_redlines': len(llm_redlines),
                'redlines': redline_models,
                'output_path': str(output_path),
                'llm_stats': self.llm_orchestrator.get_stats()
            }

            # Save result
            result_path = self.storage_path / "working" / f"{job_id}_result.json"
            result_path.write_text(json.dumps(result, default=str), encoding='utf-8')

            return result

        except Exception as e:
            print(f"Error processing document {job_id}: {e}")
            import traceback
            traceback.print_exc()

            await self._update_status(
                status_callback,
                JobStatus.ERROR,
                0,
                error_message=str(e)
            )

            return {
                'job_id': job_id,
                'status': JobStatus.ERROR,
                'error': str(e)
            }

    async def _generate_redlined_doc(
        self,
        job_id: str,
        doc: Document,
        indexer: WorkingTextIndexer,
        redlines: List[Dict]
    ) -> Path:
        """Generate Word document with track changes"""
        # Create track changes engine
        engine = TrackChangesEngine(author="ndaOK")

        # Enable track changes in document
        engine.enable_track_changes(doc)

        # Apply all redlines
        applied_count = engine.apply_all_redlines(doc, indexer, redlines)

        print(f"Job {job_id}: Applied {applied_count} of {len(redlines)} redlines")

        # Save output
        output_path = self.storage_path / "exports" / f"{job_id}_redlined.docx"
        output_path.parent.mkdir(parents=True, exist_ok=True)

        doc.save(str(output_path))

        return output_path

    async def _update_status(
        self,
        callback,
        status: JobStatus,
        progress: float,
        error_message: str = None
    ):
        """Update job status via callback"""
        if callback:
            await callback({
                'status': status,
                'progress': progress,
                'error_message': error_message,
                'updated_at': datetime.now().isoformat()
            })

    def _convert_to_models(self, redlines: List[Dict]) -> List[Dict]:
        """Convert redline dicts to serializable format"""
        models = []

        for redline in redlines:
            clause_type = redline.get('clause_type', 'unknown')

            # Get full checklist rule explanation
            rule_info = get_rule_explanation(clause_type)

            model = {
                'id': str(uuid.uuid4()),
                'rule_id': redline.get('rule_id'),
                'clause_type': clause_type,
                'start': redline['start'],
                'end': redline['end'],
                'original_text': redline['original_text'],
                'revised_text': redline.get('revised_text', ''),
                'severity': redline.get('severity', 'moderate'),
                'confidence': redline.get('confidence', 100),
                'source': redline.get('source', 'rule'),
                'explanation': redline.get('explanation') or redline.get('reasoning'),
                'validated': redline.get('validated', False),
                'user_decision': None,
                # Add full checklist rule information for UI
                'checklist_rule': {
                    'title': rule_info.get('title'),
                    'requirement': rule_info.get('requirement'),
                    'description': rule_info.get('description'),
                    'why': rule_info.get('why'),
                    'standard_language': rule_info.get('standard_language')
                }
            }

            models.append(model)

        return models


class JobQueue:
    """Simple in-memory job queue (replace with Redis for production)"""

    def __init__(self):
        self.jobs: Dict[str, Dict] = {}
        self.processor = DocumentProcessor()

    async def submit_job(self, job_id: str, file_path: str, filename: str) -> Dict:
        """Submit a new job to the queue"""
        job_info = {
            'job_id': job_id,
            'filename': filename,
            'file_path': file_path,
            'status': JobStatus.QUEUED,
            'progress': 0,
            'created_at': datetime.now(),
            'updated_at': datetime.now(),
            'result': None
        }

        self.jobs[job_id] = job_info

        # Process in background
        asyncio.create_task(self._process_job(job_id, file_path))

        return job_info

    async def _process_job(self, job_id: str, file_path: str):
        """Process job and update status"""
        async def status_callback(update):
            if job_id in self.jobs:
                self.jobs[job_id].update(update)
                self.jobs[job_id]['updated_at'] = datetime.now()

        result = await self.processor.process_document(
            job_id,
            file_path,
            status_callback
        )

        if job_id in self.jobs:
            self.jobs[job_id]['result'] = result
            self.jobs[job_id]['status'] = result.get('status', JobStatus.COMPLETE)

    def get_job_status(self, job_id: str) -> Optional[Dict]:
        """Get current job status"""
        return self.jobs.get(job_id)

    def update_redline_decision(self, job_id: str, redline_id: str, decision: str):
        """Update user decision on a redline"""
        if job_id in self.jobs and self.jobs[job_id]['result']:
            redlines = self.jobs[job_id]['result'].get('redlines', [])
            for redline in redlines:
                if redline['id'] == redline_id:
                    redline['user_decision'] = decision
                    return True
        return False

    async def export_final_document(self, job_id: str) -> Optional[Path]:
        """Export final document with only accepted changes"""
        if job_id not in self.jobs:
            return None

        job = self.jobs[job_id]
        result = job.get('result')

        if not result:
            return None

        # Filter to accepted redlines
        all_redlines = result.get('redlines', [])
        accepted = [r for r in all_redlines if r.get('user_decision') == 'accept']

        # Re-generate with only accepted changes
        file_path = job['file_path']
        doc = Document(file_path)

        indexer = WorkingTextIndexer()
        indexer.build_index(doc)

        # Convert back to dict format
        accepted_dicts = [{
            'start': r['start'],
            'end': r['end'],
            'original_text': r['original_text'],
            'revised_text': r['revised_text']
        } for r in accepted]

        engine = TrackChangesEngine(author="ndaOK")
        engine.enable_track_changes(doc)
        engine.apply_all_redlines(doc, indexer, accepted_dicts)

        output_path = Path("./storage/exports") / f"{job_id}_final.docx"
        doc.save(str(output_path))

        return output_path


class DistributedJobQueue:
    """
    Distributed job queue using Redis for horizontal scaling.
    Provides backward compatibility with existing API while adding new features.
    """

    def __init__(self, use_redis: bool = None):
        """
        Initialize job queue (Redis or in-memory based on config).

        Args:
            use_redis: Force Redis usage (None = auto-detect from env)
        """
        self.processor = DocumentProcessor()

        # Determine if Redis should be used
        if use_redis is None:
            use_redis = os.getenv("USE_REDIS_QUEUE", "false").lower() == "true"

        self.use_redis = use_redis and os.getenv("REDIS_URL") is not None

        if self.use_redis:
            self.redis_queue = RedisJobQueue(redis_url=os.getenv("REDIS_URL"))
            self.jobs = {}  # Local cache for backward compatibility
            asyncio.create_task(self._init_redis())
            asyncio.create_task(self._start_worker())
        else:
            # Fallback to in-memory queue
            self.jobs: Dict[str, Dict] = {}
            self.redis_queue = None

    async def _init_redis(self):
        """Initialize Redis connection."""
        await self.redis_queue.connect()

    async def _start_worker(self):
        """Start Redis job consumer in background."""
        if self.redis_queue:
            asyncio.create_task(
                self.redis_queue.consume_jobs(self._process_redis_job)
            )

    async def _process_redis_job(self, job_id: str, job_data: Dict) -> Dict:
        """
        Process job from Redis queue.

        Args:
            job_id: Job identifier
            job_data: Job payload

        Returns:
            Processing result
        """
        file_path = job_data['file_path']

        # Process document
        result = await self.processor.process_document(
            job_id,
            file_path,
            None  # Status updates handled via Redis
        )

        return result

    async def submit_job(
        self,
        job_id: str,
        file_path: str,
        filename: str,
        priority: str = "standard"
    ) -> Dict:
        """
        Submit a new job to the queue.

        Args:
            job_id: Unique job identifier
            file_path: Path to document
            filename: Original filename
            priority: Job priority level

        Returns:
            Job information
        """
        job_info = {
            'job_id': job_id,
            'filename': filename,
            'file_path': file_path,
            'status': JobStatus.QUEUED,
            'progress': 0,
            'created_at': datetime.now(),
            'updated_at': datetime.now(),
            'result': None
        }

        if self.use_redis:
            # Submit to Redis queue
            priority_level = {
                'low': JobPriority.LOW,
                'standard': JobPriority.STANDARD,
                'expedited': JobPriority.EXPEDITED,
                'critical': JobPriority.CRITICAL
            }.get(priority.lower(), JobPriority.STANDARD)

            await self.redis_queue.submit_job(
                {
                    'file_path': file_path,
                    'filename': filename
                },
                priority=priority_level,
                job_id=job_id
            )

            # Cache locally for backward compatibility
            self.jobs[job_id] = job_info

        else:
            # Use in-memory queue
            self.jobs[job_id] = job_info

            # Process in background
            asyncio.create_task(self._process_job(job_id, file_path))

        return job_info

    async def _process_job(self, job_id: str, file_path: str):
        """Process job in-memory (fallback mode)."""
        async def status_callback(update):
            if job_id in self.jobs:
                self.jobs[job_id].update(update)
                self.jobs[job_id]['updated_at'] = datetime.now()

        result = await self.processor.process_document(
            job_id,
            file_path,
            status_callback
        )

        if job_id in self.jobs:
            self.jobs[job_id]['result'] = result
            self.jobs[job_id]['status'] = result.get('status', JobStatus.COMPLETE)

    async def get_job_status(self, job_id: str) -> Optional[Dict]:
        """
        Get current job status.

        Args:
            job_id: Job identifier

        Returns:
            Job status information
        """
        if self.use_redis:
            # Get from Redis
            redis_status = await self.redis_queue.get_job_status(job_id)

            if redis_status:
                # Map Redis status to internal format
                status_map = {
                    RedisJobStatus.QUEUED.value: JobStatus.QUEUED,
                    RedisJobStatus.PROCESSING.value: JobStatus.PROCESSING,
                    RedisJobStatus.COMPLETED.value: JobStatus.COMPLETE,
                    RedisJobStatus.FAILED.value: JobStatus.ERROR,
                    RedisJobStatus.RETRYING.value: JobStatus.PROCESSING
                }

                return {
                    'job_id': job_id,
                    'status': status_map.get(redis_status['status'], JobStatus.QUEUED),
                    'progress': float(redis_status.get('progress', 0)),
                    'created_at': redis_status.get('created_at'),
                    'updated_at': redis_status.get('updated_at'),
                    'result': redis_status.get('result'),
                    'error': redis_status.get('error')
                }

            # Check local cache
            return self.jobs.get(job_id)

        else:
            # In-memory mode
            return self.jobs.get(job_id)

    def update_redline_decision(self, job_id: str, redline_id: str, decision: str):
        """Update user decision on a redline."""
        if job_id in self.jobs and self.jobs[job_id]['result']:
            redlines = self.jobs[job_id]['result'].get('redlines', [])
            for redline in redlines:
                if redline['id'] == redline_id:
                    redline['user_decision'] = decision
                    return True
        return False

    async def export_final_document(self, job_id: str) -> Optional[Path]:
        """Export final document with only accepted changes."""
        # Get job info
        job_info = await self.get_job_status(job_id)

        if not job_info or not job_info.get('result'):
            return None

        result = job_info['result']

        # Filter to accepted redlines
        all_redlines = result.get('redlines', [])
        accepted = [r for r in all_redlines if r.get('user_decision') == 'accept']

        # Re-generate with only accepted changes
        file_path = self.jobs[job_id]['file_path'] if job_id in self.jobs else None

        if not file_path:
            return None

        doc = Document(file_path)

        indexer = WorkingTextIndexer()
        indexer.build_index(doc)

        # Convert back to dict format
        accepted_dicts = [{
            'start': r['start'],
            'end': r['end'],
            'original_text': r['original_text'],
            'revised_text': r['revised_text']
        } for r in accepted]

        engine = TrackChangesEngine(author="ndaOK")
        engine.enable_track_changes(doc)
        engine.apply_all_redlines(doc, indexer, accepted_dicts)

        output_path = Path("./storage/exports") / f"{job_id}_final.docx"
        doc.save(str(output_path))

        return output_path

    async def get_queue_stats(self) -> Dict:
        """Get queue statistics."""
        if self.use_redis:
            return await self.redis_queue.get_queue_stats()
        else:
            return {
                'total_jobs': len(self.jobs),
                'queued': sum(1 for j in self.jobs.values() if j['status'] == JobStatus.QUEUED),
                'processing': sum(1 for j in self.jobs.values() if j['status'] == JobStatus.PROCESSING),
                'completed': sum(1 for j in self.jobs.values() if j['status'] == JobStatus.COMPLETE),
                'failed': sum(1 for j in self.jobs.values() if j['status'] == JobStatus.ERROR)
            }


# Global queue instance - automatically uses Redis if configured
job_queue = DistributedJobQueue()
