#!/usr/bin/env python3
"""
Production Upload Fix Verification Script
Tests the upload functionality after environment variable fixes
"""

import requests
import time
import sys
from pathlib import Path

# Configuration
FRONTEND_URL = "https://edgetoolspro.com"
BACKEND_URL = "https://lucky-spirit-production.up.railway.app"  # Update this with actual Railway URL
TEST_FILE = "test_nda.docx"

def test_backend_health():
    """Test if backend is accessible"""
    print("\n1. Testing Backend Health...")
    try:
        response = requests.get(f"{BACKEND_URL}/", timeout=10)
        if response.status_code == 200:
            print(f"   ✓ Backend is healthy: {response.json()}")
            return True
        else:
            print(f"   ✗ Backend returned status {response.status_code}")
            return False
    except Exception as e:
        print(f"   ✗ Backend connection failed: {e}")
        return False

def test_cors_headers():
    """Test CORS configuration"""
    print("\n2. Testing CORS Headers...")
    headers = {
        "Origin": "https://edgetoolspro.com",
        "Access-Control-Request-Method": "POST",
        "Access-Control-Request-Headers": "content-type"
    }

    try:
        # Send OPTIONS request to test CORS
        response = requests.options(
            f"{BACKEND_URL}/api/upload",
            headers=headers,
            timeout=10
        )

        cors_origin = response.headers.get("Access-Control-Allow-Origin")
        if cors_origin:
            print(f"   ✓ CORS configured: Allow-Origin = {cors_origin}")
            return True
        else:
            print(f"   ✗ No CORS headers in response")
            print(f"   Response headers: {dict(response.headers)}")
            return False
    except Exception as e:
        print(f"   ✗ CORS test failed: {e}")
        return False

def test_file_upload():
    """Test actual file upload"""
    print("\n3. Testing File Upload...")

    # Check if test file exists
    if not Path(TEST_FILE).exists():
        print(f"   ⚠ Test file '{TEST_FILE}' not found. Creating one...")
        create_test_file()

    try:
        with open(TEST_FILE, "rb") as f:
            files = {"file": (TEST_FILE, f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
            headers = {"Origin": "https://edgetoolspro.com"}

            response = requests.post(
                f"{BACKEND_URL}/api/upload",
                files=files,
                headers=headers,
                timeout=30
            )

            if response.status_code == 200:
                job_data = response.json()
                print(f"   ✓ Upload successful!")
                print(f"   Job ID: {job_data.get('job_id')}")
                print(f"   Status: {job_data.get('status')}")
                return job_data.get('job_id')
            else:
                print(f"   ✗ Upload failed with status {response.status_code}")
                print(f"   Response: {response.text}")
                return None
    except Exception as e:
        print(f"   ✗ Upload failed: {e}")
        return None

def test_job_status(job_id):
    """Test job status endpoint"""
    print(f"\n4. Testing Job Status for {job_id}...")

    try:
        response = requests.get(
            f"{BACKEND_URL}/api/jobs/{job_id}/status",
            timeout=10
        )

        if response.status_code == 200:
            status_data = response.json()
            print(f"   ✓ Status retrieved successfully")
            print(f"   Job Status: {status_data.get('status')}")
            print(f"   Progress: {status_data.get('progress')}%")
            return True
        else:
            print(f"   ✗ Status check failed with status {response.status_code}")
            return False
    except Exception as e:
        print(f"   ✗ Status check failed: {e}")
        return False

def create_test_file():
    """Create a simple test DOCX file"""
    try:
        from docx import Document
        doc = Document()
        doc.add_heading("Test NDA Document", 0)
        doc.add_paragraph("This is a test document for upload verification.")
        doc.add_paragraph("1. Test clause one")
        doc.add_paragraph("2. Test clause two")
        doc.save(TEST_FILE)
        print(f"   ✓ Created test file: {TEST_FILE}")
    except ImportError:
        print("   ✗ python-docx not installed. Install with: pip install python-docx")
        sys.exit(1)

def main():
    print("="*60)
    print("NDA REDLINE TOOL - PRODUCTION FIX VERIFICATION")
    print("="*60)
    print(f"Frontend URL: {FRONTEND_URL}")
    print(f"Backend URL:  {BACKEND_URL}")

    # Run tests
    results = {
        "Backend Health": test_backend_health(),
        "CORS Configuration": test_cors_headers(),
    }

    # Only test upload if backend is healthy
    if results["Backend Health"]:
        job_id = test_file_upload()
        results["File Upload"] = job_id is not None

        if job_id:
            time.sleep(2)  # Wait a bit for processing to start
            results["Job Status"] = test_job_status(job_id)

    # Summary
    print("\n" + "="*60)
    print("TEST RESULTS SUMMARY")
    print("="*60)

    for test_name, passed in results.items():
        status = "✓ PASSED" if passed else "✗ FAILED"
        print(f"{test_name:.<30} {status}")

    # Overall result
    all_passed = all(results.values())
    print("\n" + "="*60)
    if all_passed:
        print("✓ ALL TESTS PASSED - Upload fix is working!")
    else:
        print("✗ SOME TESTS FAILED - Please check configuration")
        print("\nTroubleshooting:")
        print("1. Verify NEXT_PUBLIC_API_URL in Vercel is set to Railway backend URL")
        print("2. Verify CORS_ORIGINS in Railway includes 'https://edgetoolspro.com'")
        print("3. Ensure both services have redeployed after changes")
        print("4. Check Railway logs for any errors")
    print("="*60)

    return 0 if all_passed else 1

if __name__ == "__main__":
    sys.exit(main())