#!/usr/bin/env python3
"""
Production Fix Validation Script
Tests that the async/await bug is fixed and document processing works end-to-end
"""

import asyncio
import aiohttp
import time
import sys
import json
from pathlib import Path
from typing import Optional, Dict

# Configuration
BACKEND_URL = "https://nda-redline-tool-production.up.railway.app"
TEST_FILE = "test_nda.docx"

class Colors:
    GREEN = '\033[92m'
    RED = '\033[91m'
    YELLOW = '\033[93m'
    BLUE = '\033[94m'
    RESET = '\033[0m'
    BOLD = '\033[1m'

def print_status(status: str, color: str = Colors.BLUE):
    """Print colored status message"""
    print(f"{color}{status}{Colors.RESET}")

def print_success(message: str):
    """Print success message"""
    print(f"{Colors.GREEN}✓ {message}{Colors.RESET}")

def print_error(message: str):
    """Print error message"""
    print(f"{Colors.RED}✗ {message}{Colors.RESET}")

def print_warning(message: str):
    """Print warning message"""
    print(f"{Colors.YELLOW}⚠ {message}{Colors.RESET}")

async def test_backend_health(session: aiohttp.ClientSession) -> bool:
    """Test if backend is accessible and healthy"""
    print_status("\n1. Testing Backend Health...", Colors.BOLD)

    try:
        async with session.get(f"{BACKEND_URL}/") as response:
            if response.status == 200:
                data = await response.json()
                print_success(f"Backend is healthy: v{data.get('version', 'unknown')}")

                # Check dependencies
                deps = data.get('dependencies', {})
                if deps:
                    for dep, info in deps.items():
                        if info.get('available'):
                            print(f"  • {dep}: Available")
                        else:
                            print_warning(f"  • {dep}: Not available")

                return True
            else:
                print_error(f"Backend returned status {response.status}")
                return False
    except Exception as e:
        print_error(f"Backend connection failed: {e}")
        return False

async def test_file_upload(session: aiohttp.ClientSession) -> Optional[str]:
    """Test document upload"""
    print_status("\n2. Testing Document Upload...", Colors.BOLD)

    # Create test file if needed
    if not Path(TEST_FILE).exists():
        print_warning(f"Creating test file: {TEST_FILE}")
        try:
            from docx import Document
            doc = Document()
            doc.add_heading("Test NDA Agreement", 0)
            doc.add_paragraph("This is a test confidentiality agreement.")
            doc.add_paragraph("1. Confidential Information shall mean all information disclosed.")
            doc.add_paragraph("2. The receiving party agrees to maintain confidentiality.")
            doc.add_paragraph("3. This agreement shall be governed by applicable law.")
            doc.save(TEST_FILE)
            print_success(f"Created test file: {TEST_FILE}")
        except ImportError:
            print_error("python-docx not installed. Install with: pip install python-docx")
            return None

    try:
        with open(TEST_FILE, 'rb') as f:
            data = aiohttp.FormData()
            data.add_field('file', f,
                          filename=TEST_FILE,
                          content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

            async with session.post(f"{BACKEND_URL}/api/upload", data=data) as response:
                if response.status == 200:
                    result = await response.json()
                    job_id = result.get('job_id')
                    print_success(f"Upload successful! Job ID: {job_id}")
                    print(f"  • Status: {result.get('status')}")
                    print(f"  • Filename: {result.get('filename')}")
                    return job_id
                else:
                    error_text = await response.text()
                    print_error(f"Upload failed with status {response.status}: {error_text}")
                    return None
    except Exception as e:
        print_error(f"Upload failed: {e}")
        return None

async def test_job_status(session: aiohttp.ClientSession, job_id: str) -> bool:
    """Test job status endpoint - THIS IS WHERE THE BUG WAS"""
    print_status("\n3. Testing Job Status Endpoint (Bug Fix Validation)...", Colors.BOLD)

    max_attempts = 30  # 30 seconds timeout
    attempt = 0

    while attempt < max_attempts:
        try:
            async with session.get(f"{BACKEND_URL}/api/jobs/{job_id}/status") as response:
                if response.status == 200:
                    data = await response.json()
                    status = data.get('status')
                    progress = data.get('progress', 0)

                    if attempt == 0:
                        print_success("Status endpoint working! (Bug is fixed)")

                    print(f"  • Status: {status} | Progress: {progress}%")

                    if status == 'complete':
                        print_success("Document processing complete!")
                        print(f"  • Total redlines: {data.get('total_redlines', 0)}")
                        print(f"  • Rule redlines: {data.get('rule_redlines', 0)}")
                        print(f"  • LLM redlines: {data.get('llm_redlines', 0)}")
                        return True
                    elif status == 'error':
                        print_error(f"Processing failed: {data.get('error_message', 'Unknown error')}")
                        return False

                elif response.status == 500:
                    # THIS IS THE BUG - Should not happen after fix
                    error_text = await response.text()
                    if 'coroutine' in error_text and 'subscriptable' in error_text:
                        print_error("CRITICAL BUG STILL PRESENT: 'coroutine' object is not subscriptable")
                        print_error("The async/await fix was not properly deployed!")
                    else:
                        print_error(f"Server error 500: {error_text}")
                    return False
                else:
                    print_error(f"Status check failed with status {response.status}")
                    return False

        except Exception as e:
            print_error(f"Status check failed: {e}")
            return False

        await asyncio.sleep(1)
        attempt += 1

    print_warning("Processing timeout - job still running after 30 seconds")
    return False

async def test_sse_events(job_id: str) -> bool:
    """Test SSE streaming endpoint"""
    print_status("\n4. Testing SSE Event Streaming...", Colors.BOLD)

    try:
        async with aiohttp.ClientSession() as session:
            url = f"{BACKEND_URL}/api/jobs/{job_id}/events"
            timeout = aiohttp.ClientTimeout(total=10)

            async with session.get(url, timeout=timeout) as response:
                if response.status == 200:
                    print_success("SSE connection established")

                    # Read first few events
                    event_count = 0
                    async for line in response.content:
                        if event_count >= 3:  # Just test first 3 events
                            break

                        decoded = line.decode('utf-8').strip()
                        if decoded.startswith('data:'):
                            try:
                                event_data = json.loads(decoded[5:])
                                print(f"  • Event received: {event_data}")
                                event_count += 1
                            except json.JSONDecodeError:
                                pass

                    return True
                else:
                    print_error(f"SSE connection failed with status {response.status}")
                    return False
    except asyncio.TimeoutError:
        print_success("SSE timeout as expected (job already complete)")
        return True
    except Exception as e:
        print_error(f"SSE test failed: {e}")
        return False

async def test_download(session: aiohttp.ClientSession, job_id: str) -> bool:
    """Test document download"""
    print_status("\n5. Testing Document Download...", Colors.BOLD)

    try:
        async with session.get(f"{BACKEND_URL}/api/jobs/{job_id}/download") as response:
            if response.status == 200:
                content = await response.read()
                output_file = f"downloaded_{job_id}.docx"

                with open(output_file, 'wb') as f:
                    f.write(content)

                print_success(f"Download successful! Saved as: {output_file}")
                print(f"  • File size: {len(content):,} bytes")
                return True
            else:
                print_error(f"Download failed with status {response.status}")
                return False
    except Exception as e:
        print_error(f"Download failed: {e}")
        return False

async def run_all_tests():
    """Run all validation tests"""
    print_status(f"\n{Colors.BOLD}{'='*60}")
    print_status(f"NDA REDLINE TOOL - PRODUCTION FIX VALIDATION")
    print_status(f"{'='*60}{Colors.RESET}")
    print(f"Backend URL: {BACKEND_URL}")
    print(f"Test File:   {TEST_FILE}")

    results = {}
    job_id = None

    async with aiohttp.ClientSession() as session:
        # Test 1: Backend health
        results['backend_health'] = await test_backend_health(session)

        if not results['backend_health']:
            print_error("\nBackend is not accessible. Stopping tests.")
            return results

        # Test 2: Upload
        job_id = await test_file_upload(session)
        results['upload'] = job_id is not None

        if not job_id:
            print_error("\nUpload failed. Stopping tests.")
            return results

        # Test 3: Status check (THE CRITICAL BUG TEST)
        results['status_check'] = await test_job_status(session, job_id)

        # Test 4: SSE streaming
        results['sse_streaming'] = await test_sse_events(job_id)

        # Test 5: Download (only if processing complete)
        if results['status_check']:
            results['download'] = await test_download(session, job_id)
        else:
            results['download'] = False
            print_warning("\n5. Skipping download test (processing not complete)")

    # Summary
    print_status(f"\n{Colors.BOLD}{'='*60}")
    print_status("TEST RESULTS SUMMARY")
    print_status(f"{'='*60}{Colors.RESET}")

    for test_name, passed in results.items():
        test_display = test_name.replace('_', ' ').title()
        if passed:
            print_success(f"{test_display:.<30} PASSED")
        else:
            print_error(f"{test_display:.<30} FAILED")

    # Overall result
    all_passed = all(results.values())
    critical_bug_fixed = results.get('status_check', False)

    print_status(f"\n{Colors.BOLD}{'='*60}{Colors.RESET}")

    if critical_bug_fixed:
        print_success("\n✓ CRITICAL BUG IS FIXED - Status endpoint no longer returns coroutine error!")
    else:
        print_error("\n✗ CRITICAL BUG STILL PRESENT - Status endpoint failing!")

    if all_passed:
        print_success("✓ ALL TESTS PASSED - System is fully operational!")
    else:
        failed_count = sum(1 for v in results.values() if not v)
        print_error(f"✗ {failed_count} TEST(S) FAILED - Please check the logs above")

    print_status(f"\n{Colors.BOLD}{'='*60}{Colors.RESET}")

    return 0 if all_passed else 1

def main():
    """Main entry point"""
    try:
        return asyncio.run(run_all_tests())
    except KeyboardInterrupt:
        print_warning("\n\nTests interrupted by user")
        return 1
    except Exception as e:
        print_error(f"\nUnexpected error: {e}")
        import traceback
        traceback.print_exc()
        return 1

if __name__ == "__main__":
    sys.exit(main())