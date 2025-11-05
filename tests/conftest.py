"""
Pytest configuration and shared fixtures for NDA Redline Tool tests
"""
import os
import sys
import pytest
from pathlib import Path
from typing import Generator
from unittest.mock import Mock, MagicMock, patch

# Add project root to Python path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

# Set test environment variables
os.environ["LOG_LEVEL"] = "DEBUG"
os.environ["ENVIRONMENT"] = "test"
os.environ["OPENAI_API_KEY"] = "sk-test-fake-key-for-testing"
os.environ["ANTHROPIC_API_KEY"] = "sk-ant-test-fake-key-for-testing"


# ============================================================================
# FastAPI Test Client Fixtures
# ============================================================================

@pytest.fixture
def test_client():
    """Provides a FastAPI test client"""
    from fastapi.testclient import TestClient
    from backend.app.main import app

    with TestClient(app) as client:
        yield client


@pytest.fixture
def async_test_client():
    """Provides an async FastAPI test client"""
    from httpx import AsyncClient
    from backend.app.main import app

    async def _client():
        async with AsyncClient(app=app, base_url="http://test") as ac:
            yield ac

    return _client


# ============================================================================
# Mock LLM Clients
# ============================================================================

@pytest.fixture
def mock_openai_client():
    """Mock OpenAI client for testing"""
    mock_client = MagicMock()

    # Mock GPT response
    mock_response = MagicMock()
    mock_response.choices = [
        MagicMock(
            message=MagicMock(
                content='{"violations": [{"clause_type": "confidentiality_term", "original_text": "two years", "revised_text": "eighteen months", "severity": "critical", "confidence": 95, "explanation": "Test violation"}]}'
            )
        )
    ]
    mock_client.chat.completions.create.return_value = mock_response

    return mock_client


@pytest.fixture
def mock_anthropic_client():
    """Mock Anthropic client for testing"""
    mock_client = MagicMock()

    # Mock Claude response
    mock_response = MagicMock()
    mock_response.content = [
        MagicMock(
            text='{"verdict": "confirm", "explanation": "This is a valid violation"}'
        )
    ]
    mock_client.messages.create.return_value = mock_response

    return mock_client


@pytest.fixture
def mock_llm_orchestrator(mock_openai_client, mock_anthropic_client):
    """Mock LLM Orchestrator with both clients"""
    with patch('backend.app.core.llm_orchestrator.LLMOrchestrator') as mock:
        instance = MagicMock()
        instance.openai_client = mock_openai_client
        instance.anthropic_client = mock_anthropic_client
        instance.analyze.return_value = {
            'violations': [],
            'total_violations': 0,
            'gpt_violations': 0,
            'claude_validations': 0
        }
        mock.return_value = instance
        yield instance


# ============================================================================
# Test Document Fixtures
# ============================================================================

@pytest.fixture
def sample_nda_text():
    """Sample NDA text for testing"""
    return """
NON-DISCLOSURE AGREEMENT

This Agreement is entered into as of January 1, 2024.

1. CONFIDENTIAL INFORMATION
The Receiving Party agrees to hold in confidence all Confidential Information
for a period of two (2) years from the date of disclosure.

2. TERM
This Agreement shall remain in effect for a period of three (3) years.

3. GOVERNING LAW
This Agreement shall be governed by the laws of the State of California.
"""


@pytest.fixture
def sample_docx_file(tmp_path):
    """Create a temporary .docx file for testing"""
    from docx import Document

    doc_path = tmp_path / "test_nda.docx"
    doc = Document()
    doc.add_paragraph("NON-DISCLOSURE AGREEMENT")
    doc.add_paragraph("This Agreement is entered into as of January 1, 2024.")
    doc.add_paragraph("Confidential Information shall be held for two (2) years.")
    doc.save(doc_path)

    return doc_path


@pytest.fixture
def sample_violation():
    """Sample violation object for testing"""
    return {
        'id': 'test-violation-1',
        'clause_type': 'confidentiality_term',
        'start': 100,
        'end': 115,
        'original_text': 'two (2) years',
        'revised_text': 'eighteen (18) months',
        'severity': 'critical',
        'confidence': 100,
        'source': 'rule',
        'explanation': 'Confidentiality term exceeds 18-month limit',
        'validated': False
    }


@pytest.fixture
def sample_redlines():
    """Sample list of redlines for testing"""
    return [
        {
            'id': 'r1',
            'clause_type': 'confidentiality_term',
            'original_text': 'two (2) years',
            'revised_text': 'eighteen (18) months',
            'severity': 'critical',
            'confidence': 100,
            'source': 'rule'
        },
        {
            'id': 'r2',
            'clause_type': 'governing_law',
            'original_text': 'State of California',
            'revised_text': 'State of Delaware',
            'severity': 'high',
            'confidence': 95,
            'source': 'gpt5'
        }
    ]


# ============================================================================
# Job Queue Fixtures
# ============================================================================

@pytest.fixture
def mock_job_queue():
    """Mock job queue for testing"""
    mock_queue = MagicMock()
    mock_queue.jobs = {}

    def create_job(job_id, filename):
        mock_queue.jobs[job_id] = {
            'job_id': job_id,
            'filename': filename,
            'status': 'queued',
            'progress': 0,
            'created_at': '2024-01-01T00:00:00',
            'updated_at': '2024-01-01T00:00:00'
        }
        return mock_queue.jobs[job_id]

    mock_queue.create_job = create_job
    mock_queue.get_job = lambda job_id: mock_queue.jobs.get(job_id)

    return mock_queue


# ============================================================================
# File System Fixtures
# ============================================================================

@pytest.fixture
def temp_storage(tmp_path):
    """Temporary storage directory structure"""
    storage = tmp_path / "storage"
    uploads = storage / "uploads"
    exports = storage / "exports"
    working = storage / "working"

    uploads.mkdir(parents=True)
    exports.mkdir(parents=True)
    working.mkdir(parents=True)

    return {
        'root': storage,
        'uploads': uploads,
        'exports': exports,
        'working': working
    }


# ============================================================================
# Rule Engine Fixtures
# ============================================================================

@pytest.fixture
def sample_rules():
    """Sample rules for testing"""
    return [
        {
            'id': 'CONF_TERM_001',
            'name': 'Confidentiality Term Limit',
            'pattern': r'(\w+\s*\(\d+\)\s*years?)',
            'action': 'replace',
            'replacement': 'eighteen (18) months',
            'severity': 'critical',
            'clause_type': 'confidentiality_term',
            'explanation': 'Replace term exceeding 18 months'
        },
        {
            'id': 'GOV_LAW_001',
            'name': 'Governing Law - Delaware',
            'pattern': r'State of (?!Delaware)\w+',
            'action': 'replace',
            'replacement': 'State of Delaware',
            'severity': 'high',
            'clause_type': 'governing_law',
            'explanation': 'Change governing law to Delaware'
        }
    ]


# ============================================================================
# Environment Fixtures
# ============================================================================

@pytest.fixture
def test_env():
    """Test environment variables"""
    original_env = os.environ.copy()

    # Set test environment
    os.environ.update({
        'ENVIRONMENT': 'test',
        'LOG_LEVEL': 'DEBUG',
        'MAX_FILE_SIZE_MB': '10',
        'VALIDATION_RATE': '0.15',
        'CONFIDENCE_THRESHOLD': '95'
    })

    yield os.environ

    # Restore original environment
    os.environ.clear()
    os.environ.update(original_env)


# ============================================================================
# Markers Configuration
# ============================================================================

def pytest_configure(config):
    """Register custom markers"""
    config.addinivalue_line("markers", "unit: Unit tests")
    config.addinivalue_line("markers", "integration: Integration tests")
    config.addinivalue_line("markers", "slow: Slow running tests")
    config.addinivalue_line("markers", "fast: Fast running tests")
    config.addinivalue_line("markers", "requires_api_keys: Tests requiring real API keys")
    config.addinivalue_line("markers", "requires_redis: Tests requiring Redis")
    config.addinivalue_line("markers", "smoke: Smoke tests")


# ============================================================================
# Test Helpers
# ============================================================================

@pytest.fixture
def assert_logs():
    """Helper for asserting log messages"""
    def _assert_logs(caplog, level, message_contains):
        """Check if a log message was recorded"""
        for record in caplog.records:
            if record.levelname == level and message_contains in record.message:
                return True
        return False
    return _assert_logs


# ============================================================================
# Cleanup
# ============================================================================

@pytest.fixture(autouse=True)
def cleanup_test_files(tmp_path):
    """Automatically cleanup test files after each test"""
    yield
    # Cleanup happens automatically with tmp_path
